#!/usr/bin/env python3
"""
Script de récupération et d'envoi de résumés d'actualités.

- Respectez les conditions d'utilisation des API ou des sites que vous scrapez.
  Note : Le scraping peut violer les TOS. Envisagez d'utiliser des APIs officielles comme NewsAPI (https://newsapi.org/).
- Stockez les informations d'identification dans un fichier `.env` ou un gestionnaire de secrets sécurisé.
- Gérez les exceptions et les problèmes réseau de manière robuste.
"""

import logging
import os
import smtplib
import time
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from urllib.parse import urljoin

import requests
from bs4 import BeautifulSoup
from docx import Document
from requests.exceptions import RequestException
from smtplib import SMTPException

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    logging.warning("python-dotenv n'est pas installé. Utilisez les variables d'environnement système ou installez python-dotenv.")

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s: %(message)s")

# Constantes
REQUEST_TIMEOUT = 10
SMTP_TIMEOUT = 20
DELAY_BETWEEN_REQUESTS = 1  # secondes

SMTP_SERVER = os.getenv("SMTP_SERVER", "smtp.gmail.com")
PORT = int(os.getenv("SMTP_PORT", 587))
EMAIL_ADDRESS = os.getenv("EMAIL_ADDRESS")
PASSWORD = os.getenv("EMAIL_PASSWORD")
RECIPIENT_EMAIL = os.getenv("RECIPIENT_EMAIL", "destinataire@example.com")
NEWS_CATEGORIES = ["innovation", "economie_standard_circulaire", "ecologie"]

if not EMAIL_ADDRESS or not PASSWORD:
    raise ValueError(
        "Les variables d'environnement EMAIL_ADDRESS et EMAIL_PASSWORD doivent être définies. "
        "Stockez-les dans un fichier .env et assurez-vous que .env est chargé."
    )


def fetch_news(category):
    """Récupère une liste d'URLs d'articles pour une catégorie donnée."""
    urls = {
        "innovation": "https://techcrunch.com",
        "economie_standard_circulaire": "https://www.eurecirc.fr/",
        "ecologie": "https://fr.environnement.etic.fr/",
    }
    url = urls.get(category)
    if not url:
        logging.warning("Catégorie inconnue : %s", category)
        return []

    try:
        response = requests.get(url, timeout=REQUEST_TIMEOUT)
        response.raise_for_status()
    except RequestException as exc:
        logging.error("Erreur réseau lors de la récupération de %s : %s", category, exc)
        return []

    soup = BeautifulSoup(response.text, "html.parser")
    articles = soup.find_all("article", class_="post-summary")
    links = []
    for article in articles:
        try:
            href = article.h2.a["href"]
            # Assurer que l'URL est absolue
            full_url = urljoin(url, href)
            links.append(full_url)
        except (AttributeError, TypeError):
            continue
    logging.info("%d articles trouvés pour la catégorie %s", len(links), category)
    return links


def scrape_article(url):
    """Récupère le titre et le contenu principal d'un article."""
    time.sleep(DELAY_BETWEEN_REQUESTS)  # Délai pour éviter d'être bloqué
    try:
        response = requests.get(url, timeout=REQUEST_TIMEOUT)
        response.raise_for_status()
    except RequestException as exc:
        logging.warning("Impossible de récupérer l'article %s : %s", url, exc)
        return {"title": "Article indisponible", "content": ""}

    soup = BeautifulSoup(response.text, "html.parser")
    title = soup.title.string if soup.title else "Sans titre"
    content = "\n".join([p.get_text(strip=True) for p in soup.find_all("p")])
    return {"title": title, "content": content}


def generate_summary():
    """Construit un résumé des articles par catégorie."""
    summary = {}
    for category in NEWS_CATEGORIES:
        urls = fetch_news(category)
        articles = [scrape_article(url) for url in urls]
        summary[category] = "\n\n".join(
            f"{article['title']}\n{article['content'][:400]}" for article in articles if article["content"]
        )
    return summary


def send_email(summary):
    """Envoie le résumé par e-mail en utilisant SMTP."""
    if not summary:
        logging.warning("Aucun résumé à envoyer.")
        return

    msg = MIMEMultipart()
    msg["From"] = EMAIL_ADDRESS
    msg["To"] = RECIPIENT_EMAIL
    msg["Subject"] = "Résumé de la Semaine - Actualités Techno, Économie et Écologie"

    body = "Voici le résumé des nouvelles pour cette semaine:\n\n"
    for category, content in summary.items():
        body += f"Catégorie: {category}\n{content}\n\n"
    msg.attach(MIMEText(body, "plain"))

    try:
        with smtplib.SMTP(SMTP_SERVER, PORT, timeout=SMTP_TIMEOUT) as server:
            server.starttls()
            server.login(EMAIL_ADDRESS, PASSWORD)
            server.sendmail(EMAIL_ADDRESS, RECIPIENT_EMAIL, msg.as_string())
        logging.info("E-mail envoyé à %s", RECIPIENT_EMAIL)
    except SMTPException as exc:
        logging.error("Erreur lors de l'envoi de l'e-mail : %s", exc)


def save_to_word(summary):
    """Enregistre le résumé dans un fichier Word."""
    doc = Document()
    for category, content in summary.items():
        doc.add_heading(f"Catégorie: {category}", level=1)
        doc.add_paragraph(content)
    doc.save("summary_weekly.docx")
    logging.info("Fichier Word enregistré : summary_weekly.docx")


def main():
    try:
        summary = generate_summary()
        send_email(summary)
        save_to_word(summary)
    except Exception as exc:
        logging.exception("Une erreur inattendue est survenue : %s", exc)


if __name__ == "__main__":
    main()
